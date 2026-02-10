from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from docx.oxml.shared import qn, OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import date
from decimal import Decimal
import re
import os
import sys

# Ruta absoluta del proyecto
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(BASE_DIR, ".."))
sys.path.append(project_root)

from db.database import fetch_invoice_data
from db.uDB import *  # Asegúrate que tengas este método en tu módulo db.uDB


# Función para reemplazar texto en el documento directamente
def replace_text_in_document(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old_text, new_text)

    # También buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(old_text, new_text)


# Parámetros de factura
RNCemisor = "106012612"
eNCF = "E310000000002"

# Conexión a la base de datos
GConfig.cargar(1)
cn1 = ConectarDB()

# Obtener datos
company_data, invoice_data, products, formas_pago, descuento_recargo = (
    fetch_invoice_data(cn1, RNCemisor, eNCF)
)

if not invoice_data:
    raise ValueError("No se pudo cargar la factura.")

# Obtener datos de la empresa
empresa = {
    "nombre_empresa": company_data["nombre_empresa"],
    "direccion": company_data["direccion"],
    "rnc": company_data["rnc"],
    "telefono": company_data["telefono"],
}

# Preparar lista de productos
productos = [
    {
        "nombre": p["descripcion"],
        "cantidad": p["cantidad"],
        "precio": f'{Decimal(p["precio"]):,.2f}',
        "total": f'{Decimal(p["importe"]):,.2f}',
    }
    for p in products
]

# Totales
subtotal = f'{Decimal(invoice_data["Monto_gravado"]):,.2f}'
iva = f'{Decimal(invoice_data["TotalITBIS"]):,.2f}'
total = f'{Decimal(invoice_data["Monto_total"]):,.2f}'
descuento = (
    f'{Decimal(descuento_recargo["monto"]):,.2f}' if descuento_recargo else "0.00"
)

# Dividir productos por página
productos_por_pagina = 15
paginas = [
    productos[i : i + productos_por_pagina]
    for i in range(0, len(productos), productos_por_pagina)
]

archivos_generados = []

# Procesar cada grupo de productos
for i, grupo in enumerate(paginas):
    es_ultima_pagina = i == len(paginas) - 1
    plantilla_nombre = "factura_parcial.docx"

    plantilla_ruta = os.path.join(BASE_DIR, plantilla_nombre)

    if not os.path.exists(plantilla_ruta):
        raise FileNotFoundError(f"Plantilla no encontrada: {plantilla_ruta}")

    # En lugar de usar DocxTemplate, usamos Document directamente
    doc = Document(plantilla_ruta)

    # Reemplazar variables manualmente
    replace_text_in_document(doc, "{{ nombre_empresa }}", empresa["nombre_empresa"])

    # También reemplaza otros campos necesarios
    replace_text_in_document(doc, "{{ factura_id }}", str(invoice_data["numero"]))
    replace_text_in_document(
        doc,
        "{{ fecha }}",
        (
            invoice_data["fecha"].strftime("%d/%m/%Y")
            if hasattr(invoice_data["fecha"], "strftime")
            else str(invoice_data["fecha"])
        ),
    )
    replace_text_in_document(
        doc, "{{ cliente_nombre }}", invoice_data["nombre_cliente"]
    )

    # Añadir productos a la tabla (si existe)
    if len(doc.tables) > 0:
        # Asumimos que la tabla de productos es la primera tabla
        tabla_productos = doc.tables[0]
        # Buscar la fila de plantilla (puede variar según la estructura del documento)
        fila_plantilla = None
        for i, row in enumerate(tabla_productos.rows):
            if any("{{ producto" in cell.text for cell in row.cells):
                fila_plantilla = row
                fila_plantilla_index = i
                break

        if fila_plantilla:
            # Eliminar la fila de plantilla
            tabla_productos._tbl.remove(fila_plantilla._tr)

            # Agregar filas de productos
            for producto in grupo:
                nueva_fila = tabla_productos.add_row()
                if len(nueva_fila.cells) >= 4:  # Asumiendo que hay 4 columnas
                    nueva_fila.cells[0].text = producto["nombre"]
                    nueva_fila.cells[1].text = str(producto["cantidad"])
                    nueva_fila.cells[2].text = producto["precio"]
                    nueva_fila.cells[3].text = producto["total"]

    # Reemplazar más variables si es la última página
    if es_ultima_pagina:
        replace_text_in_document(doc, "{{ subtotal }}", subtotal)
        replace_text_in_document(doc, "{{ iva }}", iva)
        replace_text_in_document(doc, "{{ total }}", total)
        replace_text_in_document(doc, "{{ descuento }}", descuento)
        replace_text_in_document(
            doc, "{{ telefono_cliente }}", invoice_data.get("telefono_cliente", "")
        )
        replace_text_in_document(
            doc, "{{ direccion_cliente }}", invoice_data.get("direccion_cliente", "")
        )
        replace_text_in_document(doc, "{{ ncf }}", invoice_data.get("ncf", ""))
        replace_text_in_document(
            doc, "{{ vendedor }}", invoice_data.get("nombre_vendedor", "")
        )
        replace_text_in_document(
            doc, "{{ observacion }}", invoice_data.get("observacion", "")
        )
        replace_text_in_document(doc, "{{ cajero }}", invoice_data.get("cajero", ""))

    nombre_archivo = os.path.join(BASE_DIR, f"pagina_{i+1}.docx")
    doc.save(nombre_archivo)
    archivos_generados.append(nombre_archivo)

# Unir páginas en un solo documento
documento_principal = Document(archivos_generados[0])
composer = Composer(documento_principal)

for archivo in archivos_generados[1:]:
    composer.append(Document(archivo))

output_final = os.path.join(BASE_DIR, "factura_final_multipagina.docx")
composer.save(output_final)

# Limpiar archivos temporales
for archivo in archivos_generados:
    os.remove(archivo)

print("Factura generada correctamente:", output_final)
